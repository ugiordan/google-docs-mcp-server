"""Convert markdown to .docx using python-docx.

Bypasses the Google Docs batchUpdate API by generating a .docx file in memory.
Google Drive's native .docx import handles tables, code blocks, and formatting
correctly, which the batchUpdate approach cannot do reliably for complex documents.
"""

import io
from urllib.parse import urlparse

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree

from mcp_server.services.markdown_converter import parse_markdown

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
_NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]

_ALLOWED_URL_SCHEMES = {"http", "https", "mailto"}


def _add_hyperlink(paragraph, url, text):
    """Add a hyperlink run to a paragraph."""
    parsed = urlparse(url)
    if not parsed.scheme or parsed.scheme.lower() not in _ALLOWED_URL_SCHEMES:
        # Render as plain text instead of a clickable link
        paragraph.add_run(text)
        return

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run_elem.append(rPr)

    run_text = OxmlElement("w:t")
    run_text.set(qn("xml:space"), "preserve")
    run_text.text = text
    run_elem.append(run_text)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_runs(paragraph, runs):
    """Add formatted runs to a paragraph."""
    for run_data in runs:
        text = run_data.get("text", "")
        if not text:
            continue

        if run_data.get("link"):
            _add_hyperlink(paragraph, run_data["link"], text)
            continue

        run = paragraph.add_run(text)
        if run_data.get("bold"):
            run.bold = True
        if run_data.get("italic"):
            run.italic = True
        if run_data.get("code"):
            run.font.name = "Roboto Mono"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 112, 0)
        if run_data.get("strikethrough"):
            run.font.strike = True


def _add_table(doc, block):
    """Add a table from a table block."""
    rows_data = block["rows"]
    if not rows_data:
        return

    num_rows = len(rows_data)
    num_cols = max(len(r) for r in rows_data)

    table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row in enumerate(rows_data):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                if block.get("has_header") and i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def _add_code_block(doc, block):
    """Add a code block with monospace font and gray shading."""
    p = doc.add_paragraph()

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)

    run = p.add_run(block["text"])
    run.font.name = "Roboto Mono"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 112, 0)


def _add_blockquote(doc, block):
    """Add a blockquote with indentation and left border."""
    p = doc.add_paragraph()
    depth = block.get("depth", 1)
    p.paragraph_format.left_indent = Pt(36 * depth)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "CCCCCC")
    pBdr.append(left)
    pPr.append(pBdr)

    runs = block.get("runs", [{"text": block.get("text", "")}])
    _add_runs(p, runs)


def _add_horizontal_rule(doc):
    """Add a horizontal rule as a paragraph with bottom border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_style_font(style, font_name):
    """Set font on a style, overriding theme font references.

    python-docx's font.name only sets w:ascii and w:hAnsi. Text that
    references the theme font (e.g. Calibri as minorHAnsi) won't pick
    up the style-level override. Setting all four rFonts attributes
    ensures the font applies regardless of theme.
    """
    style.font.name = font_name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def _override_theme_fonts(doc, body_font, heading_font=None):
    """Override the document theme's minor/major fonts.

    The default python-docx template uses Calibri/Cambria as theme fonts.
    Any text that references the theme (rather than an explicit font)
    falls back to these defaults even when styles specify a different
    font. Changing the theme XML itself eliminates this fallback.
    """
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    for rel in doc.part.rels.values():
        if "theme" not in rel.reltype:
            continue
        theme_part = rel.target_part
        try:
            theme_xml = etree.fromstring(theme_part.blob)
        except etree.XMLSyntaxError:
            continue
        minor = theme_xml.find(".//a:minorFont/a:latin", ns)
        if minor is not None:
            minor.set("typeface", body_font)
        major = theme_xml.find(".//a:majorFont/a:latin", ns)
        if major is not None:
            major.set("typeface", heading_font or body_font)
        theme_part._blob = etree.tostring(
            theme_xml, xml_declaration=True, encoding="UTF-8", standalone=True
        )


def _apply_template_styles(doc, styles):
    """Apply template styles to the document's built-in styles."""
    style_mapping = {
        "NORMAL_TEXT": "Normal",
        "HEADING_1": "Heading 1",
        "HEADING_2": "Heading 2",
        "HEADING_3": "Heading 3",
        "HEADING_4": "Heading 4",
        "HEADING_5": "Heading 5",
        "HEADING_6": "Heading 6",
    }

    body_font = None
    heading_font = None

    for style_type, doc_style_name in style_mapping.items():
        if style_type not in styles:
            continue
        props = styles[style_type]
        try:
            doc_style = doc.styles[doc_style_name]
        except KeyError:
            continue

        if "font_family" in props:
            _set_style_font(doc_style, props["font_family"])
            if style_type == "NORMAL_TEXT":
                body_font = props["font_family"]
            elif style_type == "HEADING_1" and heading_font is None:
                heading_font = props["font_family"]
        if "font_size" in props:
            doc_style.font.size = Pt(props["font_size"])
        if "foreground_color" in props:
            color_data = props["foreground_color"]
            rgb = color_data.get("rgbColor", {})
            r = round(rgb.get("red", 0) * 255)
            g = round(rgb.get("green", 0) * 255)
            b = round(rgb.get("blue", 0) * 255)
            doc_style.font.color.rgb = RGBColor(r, g, b)

    # Override the document theme to prevent Calibri fallback
    if body_font:
        _override_theme_fonts(doc, body_font, heading_font)


def blocks_to_docx(blocks: list[dict], styles: dict | None = None) -> bytes:
    """Convert parsed markdown blocks to a .docx file.

    Args:
        blocks: List of parsed markdown blocks from parse_markdown()
        styles: Optional template styles from extract_template_styles()

    Returns:
        Bytes of the generated .docx file
    """
    doc = Document()

    if styles:
        _apply_template_styles(doc, styles)

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = block.get("level", 1)
            p = doc.add_heading(level=level)
            p.clear()
            runs = block.get("runs", [{"text": block.get("text", "")}])
            _add_runs(p, runs)

        elif btype == "paragraph":
            p = doc.add_paragraph()
            runs = block.get("runs", [{"text": block.get("text", "")}])
            _add_runs(p, runs)

        elif btype == "table":
            _add_table(doc, block)

        elif btype == "code_block":
            _add_code_block(doc, block)

        elif btype == "list_item":
            ordered = block.get("ordered", False)
            nesting = block.get("nesting_level", 0)
            style_list = _NUMBER_STYLES if ordered else _BULLET_STYLES
            style_name = style_list[min(nesting, len(style_list) - 1)]
            p = doc.add_paragraph(style=style_name)
            runs = block.get("runs", [{"text": block.get("text", "")}])
            _add_runs(p, runs)

        elif btype == "blockquote":
            _add_blockquote(doc, block)

        elif btype == "horizontal_rule":
            _add_horizontal_rule(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_docx(markdown_content: str, styles: dict | None = None) -> bytes:
    """Convert markdown text to a .docx file.

    Args:
        markdown_content: Raw markdown string
        styles: Optional template styles from extract_template_styles()

    Returns:
        Bytes of the generated .docx file
    """
    blocks = parse_markdown(markdown_content)
    return blocks_to_docx(blocks, styles)
