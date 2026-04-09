"""Tests for mcp_server.services.docx_converter."""

import io

from docx import Document
from docx.shared import Pt

from mcp_server.services.docx_converter import (
    _DOCX_MIME,
    blocks_to_docx,
    markdown_to_docx,
)


def _load_doc(docx_bytes: bytes) -> Document:
    """Load a Document from bytes."""
    return Document(io.BytesIO(docx_bytes))


# ---------------------------------------------------------------------------
# blocks_to_docx tests
# ---------------------------------------------------------------------------


class TestBlocksToDocx:
    def test_empty_blocks(self):
        result = blocks_to_docx([])
        doc = _load_doc(result)
        # Should be a valid docx with just the default empty paragraph
        assert len(doc.paragraphs) >= 0

    def test_heading(self):
        blocks = [
            {
                "type": "heading",
                "level": 1,
                "text": "Title",
                "runs": [{"text": "Title"}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        assert any("Title" in p.text for p in doc.paragraphs)

    def test_heading_levels(self):
        blocks = [
            {
                "type": "heading",
                "level": i,
                "text": f"H{i}",
                "runs": [{"text": f"H{i}"}],
            }
            for i in range(1, 4)
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        texts = [p.text for p in doc.paragraphs]
        assert "H1" in texts
        assert "H2" in texts
        assert "H3" in texts

    def test_paragraph(self):
        blocks = [
            {
                "type": "paragraph",
                "text": "Hello world",
                "runs": [{"text": "Hello world"}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        assert any("Hello world" in p.text for p in doc.paragraphs)

    def test_bold_run(self):
        blocks = [
            {
                "type": "paragraph",
                "text": "normal bold",
                "runs": [{"text": "normal "}, {"text": "bold", "bold": True}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "bold" in p.text][0]
        bold_runs = [r for r in p.runs if r.bold]
        assert len(bold_runs) >= 1
        assert bold_runs[0].text == "bold"

    def test_italic_run(self):
        blocks = [
            {
                "type": "paragraph",
                "text": "normal italic",
                "runs": [{"text": "normal "}, {"text": "italic", "italic": True}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "italic" in p.text][0]
        italic_runs = [r for r in p.runs if r.italic]
        assert len(italic_runs) >= 1

    def test_code_run(self):
        blocks = [
            {
                "type": "paragraph",
                "text": "use print()",
                "runs": [{"text": "use "}, {"text": "print()", "code": True}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "print()" in p.text][0]
        code_runs = [r for r in p.runs if r.font.name == "Courier New"]
        assert len(code_runs) >= 1

    def test_strikethrough_run(self):
        blocks = [
            {
                "type": "paragraph",
                "text": "deleted",
                "runs": [{"text": "deleted", "strikethrough": True}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "deleted" in p.text][0]
        strike_runs = [r for r in p.runs if r.font.strike]
        assert len(strike_runs) >= 1

    def test_table_basic(self):
        blocks = [
            {
                "type": "table",
                "rows": [["A", "B"], ["1", "2"]],
                "has_header": True,
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert table.cell(0, 0).text == "A"
        assert table.cell(0, 1).text == "B"
        assert table.cell(1, 0).text == "1"
        assert table.cell(1, 1).text == "2"

    def test_table_header_bold(self):
        blocks = [
            {
                "type": "table",
                "rows": [["H1", "H2"], ["a", "b"]],
                "has_header": True,
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        table = doc.tables[0]
        header_cell = table.cell(0, 0)
        for p in header_cell.paragraphs:
            for r in p.runs:
                assert r.bold is True

    def test_table_no_header(self):
        blocks = [
            {
                "type": "table",
                "rows": [["x", "y"]],
                "has_header": False,
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        table = doc.tables[0]
        assert table.cell(0, 0).text == "x"

    def test_multiple_tables(self):
        blocks = [
            {"type": "table", "rows": [["A"]], "has_header": False},
            {"type": "paragraph", "text": "Middle", "runs": [{"text": "Middle"}]},
            {"type": "table", "rows": [["B"]], "has_header": False},
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        assert len(doc.tables) == 2

    def test_code_block(self):
        blocks = [{"type": "code_block", "text": "print('hello')"}]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "print" in p.text][0]
        assert p.runs[0].font.name == "Courier New"

    def test_list_item_unordered(self):
        blocks = [
            {
                "type": "list_item",
                "text": "Item one",
                "ordered": False,
                "nesting_level": 0,
                "runs": [{"text": "Item one"}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "Item one" in p.text][0]
        assert p.style.name == "List Bullet"

    def test_list_item_ordered(self):
        blocks = [
            {
                "type": "list_item",
                "text": "Step one",
                "ordered": True,
                "nesting_level": 0,
                "runs": [{"text": "Step one"}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "Step one" in p.text][0]
        assert p.style.name == "List Number"

    def test_nested_list(self):
        blocks = [
            {
                "type": "list_item",
                "text": "Nested",
                "ordered": False,
                "nesting_level": 1,
                "runs": [{"text": "Nested"}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "Nested" in p.text][0]
        assert p.style.name == "List Bullet 2"

    def test_deep_nested_list_clamps(self):
        blocks = [
            {
                "type": "list_item",
                "text": "Deep",
                "ordered": False,
                "nesting_level": 5,
                "runs": [{"text": "Deep"}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "Deep" in p.text][0]
        assert p.style.name == "List Bullet 3"

    def test_blockquote(self):
        blocks = [
            {
                "type": "blockquote",
                "text": "Quoted text",
                "depth": 1,
                "runs": [{"text": "Quoted text"}],
            }
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        p = [p for p in doc.paragraphs if "Quoted text" in p.text][0]
        assert p.paragraph_format.left_indent is not None

    def test_horizontal_rule(self):
        blocks = [{"type": "horizontal_rule"}]
        doc = _load_doc(blocks_to_docx(blocks))
        # HR creates a paragraph (possibly empty). The border is in XML.
        assert len(doc.paragraphs) >= 1

    def test_link_run(self):
        blocks = [
            {
                "type": "paragraph",
                "text": "click here",
                "runs": [
                    {"text": "click "},
                    {"text": "here", "link": "https://example.com"},
                ],
            }
        ]
        # Should not raise, and the doc should contain the text
        doc = _load_doc(blocks_to_docx(blocks))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "click" in full_text

    def test_template_styles_applied(self):
        blocks = [
            {
                "type": "heading",
                "level": 1,
                "text": "Title",
                "runs": [{"text": "Title"}],
            }
        ]
        styles = {"HEADING_1": {"font_family": "Georgia", "font_size": 28}}
        doc = _load_doc(blocks_to_docx(blocks, styles))
        h1_style = doc.styles["Heading 1"]
        assert h1_style.font.name == "Georgia"
        assert h1_style.font.size == Pt(28)

    def test_template_styles_with_color(self):
        styles = {
            "NORMAL_TEXT": {
                "foreground_color": {
                    "rgbColor": {"red": 0.2, "green": 0.4, "blue": 0.6}
                }
            }
        }
        blocks = [
            {"type": "paragraph", "text": "Colored", "runs": [{"text": "Colored"}]}
        ]
        doc = _load_doc(blocks_to_docx(blocks, styles))
        normal_style = doc.styles["Normal"]
        assert normal_style.font.color.rgb is not None

    def test_mixed_content(self):
        blocks = [
            {
                "type": "heading",
                "level": 1,
                "text": "Title",
                "runs": [{"text": "Title"}],
            },
            {"type": "paragraph", "text": "Intro", "runs": [{"text": "Intro"}]},
            {"type": "table", "rows": [["A", "B"], ["1", "2"]], "has_header": True},
            {"type": "code_block", "text": "x = 1"},
            {
                "type": "list_item",
                "text": "Item",
                "ordered": False,
                "nesting_level": 0,
                "runs": [{"text": "Item"}],
            },
            {"type": "horizontal_rule"},
            {
                "type": "blockquote",
                "text": "Quote",
                "depth": 1,
                "runs": [{"text": "Quote"}],
            },
        ]
        doc = _load_doc(blocks_to_docx(blocks))
        assert len(doc.tables) == 1
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Title" in full_text
        assert "Intro" in full_text
        assert "x = 1" in full_text

    def test_empty_table_rows(self):
        blocks = [{"type": "table", "rows": [], "has_header": False}]
        # Should not raise
        doc = _load_doc(blocks_to_docx(blocks))
        assert len(doc.tables) == 0

    def test_blocks_without_runs_key(self):
        blocks = [{"type": "paragraph", "text": "No runs key"}]
        doc = _load_doc(blocks_to_docx(blocks))
        assert any("No runs key" in p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# markdown_to_docx integration tests
# ---------------------------------------------------------------------------


class TestMarkdownToDocx:
    def test_simple_markdown(self):
        md = "# Hello\n\nWorld"
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        texts = [p.text for p in doc.paragraphs]
        assert "Hello" in texts
        assert "World" in texts

    def test_markdown_with_table(self):
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "A"
        assert doc.tables[0].cell(1, 0).text == "1"

    def test_markdown_with_code_block(self):
        md = "```python\nprint('hello')\n```"
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        code_paras = [p for p in doc.paragraphs if "print" in p.text]
        assert len(code_paras) >= 1

    def test_markdown_with_bold_and_italic(self):
        md = "This is **bold** and *italic* text"
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        p = doc.paragraphs[0]
        bold_runs = [r for r in p.runs if r.bold]
        assert len(bold_runs) >= 1

    def test_markdown_with_list(self):
        md = "- item one\n- item two"
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        list_paras = [p for p in doc.paragraphs if "item" in p.text]
        assert len(list_paras) == 2

    def test_markdown_mixed_complex(self):
        md = (
            "# Title\n\n"
            "Some **bold** text.\n\n"
            "| Col1 | Col2 |\n|---|---|\n| a | b |\n\n"
            "```\ncode\n```\n\n"
            "- bullet\n\n"
            "> quote\n\n"
            "---\n\n"
            "End."
        )
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        assert len(doc.tables) == 1
        full = " ".join(p.text for p in doc.paragraphs)
        assert "Title" in full
        assert "code" in full
        assert "End." in full

    def test_with_template_styles(self):
        md = "# Styled heading"
        styles = {"HEADING_1": {"font_family": "Times New Roman"}}
        result = markdown_to_docx(md, styles)
        doc = _load_doc(result)
        assert doc.styles["Heading 1"].font.name == "Times New Roman"

    def test_returns_valid_docx_bytes(self):
        result = markdown_to_docx("Hello")
        # .docx files start with PK (ZIP signature)
        assert result[:2] == b"PK"

    def test_docx_mime_constant(self):
        assert "wordprocessingml" in _DOCX_MIME


# ---------------------------------------------------------------------------
# Regression: complex document with many tables
# ---------------------------------------------------------------------------


class TestComplexDocument:
    def test_many_tables(self):
        parts = []
        for i in range(15):
            parts.append(f"## Section {i}\n\n")
            parts.append(f"| Header{i} | Value |\n|---|---|\n| key{i} | val{i} |\n\n")
        md = "".join(parts)
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        assert len(doc.tables) == 15

    def test_table_with_inline_formatting_in_cells(self):
        # Cells with inline code, links, etc. should not crash
        md = (
            "| `code` | [link](https://example.com) |\n|---|---|\n| normal | **bold** |"
        )
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        assert len(doc.tables) == 1

    def test_code_blocks_between_tables(self):
        md = (
            "| A |\n|---|\n| 1 |\n\n"
            "```go\nfunc main() {}\n```\n\n"
            "| B |\n|---|\n| 2 |\n\n"
        )
        result = markdown_to_docx(md)
        doc = _load_doc(result)
        assert len(doc.tables) == 2
        code_paras = [p for p in doc.paragraphs if "func main" in p.text]
        assert len(code_paras) >= 1
